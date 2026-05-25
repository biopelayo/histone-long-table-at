"""OLA 9.4 · aplica cambios v3.2.0 a CHAPTER_METHODS_rev_1604_LV.docx.

Usa python-docx 1.2.0. Lanzar con: python OLA9_4_apply_docx.py
"""
import sys
import io
import re

# UTF-8 en stdout para Windows / Git Bash
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = r"D:/AT_virgen/CHAPTER_METHODS_rev_1604_LV.docx"
doc = Document(SRC)


# ---------------------------------------------------------------------------
# (a) Reemplazo global EpiProfile_PLANTS -> EpiProfile-PLANTS
# ---------------------------------------------------------------------------
n_runs_replaced = 0


def replace_in_paragraph(paragraph, old, new):
    """Reemplazo preservando formato si es posible; fallback a primer run."""
    global n_runs_replaced
    if any(old in r.text for r in paragraph.runs):
        for r in paragraph.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                n_runs_replaced += 1
        return True
    full = "".join(r.text for r in paragraph.runs)
    if old in full:
        new_full = full.replace(old, new)
        runs = list(paragraph.runs)
        if runs:
            runs[0].text = new_full
            for r in runs[1:]:
                r.text = ""
            n_runs_replaced += 1
            return True
    return False


n_paragraphs_replaced = 0
for p in doc.paragraphs:
    if "EpiProfile_PLANTS" in p.text:
        if replace_in_paragraph(p, "EpiProfile_PLANTS", "EpiProfile-PLANTS"):
            n_paragraphs_replaced += 1
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "EpiProfile_PLANTS" in p.text:
                    if replace_in_paragraph(p, "EpiProfile_PLANTS", "EpiProfile-PLANTS"):
                        n_paragraphs_replaced += 1

print(
    f"(a) EpiProfile_PLANTS -> EpiProfile-PLANTS: "
    f"{n_paragraphs_replaced} parrafos, {n_runs_replaced} runs editados"
)


# ---------------------------------------------------------------------------
# (b) Eliminar notas embebidas del codirector LV.
# ---------------------------------------------------------------------------
patterns_to_remove = [
    re.compile(r"^\s*Notas\s*:\s*estas\s+las\s+desarrollaria", re.IGNORECASE),
    re.compile(r"^\s*LO\s+PODEMOS\s+REDUCIR\s*$", re.IGNORECASE),
    re.compile(r"^\s*Esta\s+table\s+est[aá]\s+bien", re.IGNORECASE),
    re.compile(r"^\s*Esta\s+figura\s+tambi[eé]n\s+est[aá]\s+bien", re.IGNORECASE),
    re.compile(r"^\s*E:\\EpiProfile_20_AT", re.IGNORECASE),
]

removed = []
for p in list(doc.paragraphs):
    text = p.text
    for pat in patterns_to_remove:
        if pat.search(text):
            removed.append(text[:80])
            el = p._element
            el.getparent().remove(el)
            break

print(f"(b) Notas LV eliminadas: {len(removed)}")
for r in removed:
    print(f"    - {r}")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def find_paragraph_by_substring(substr):
    for i, p in enumerate(doc.paragraphs):
        if substr in p.text:
            return i, p
    return -1, None


def insert_paragraph_after(paragraph, style=None):
    new_xml = OxmlElement("w:p")
    paragraph._element.addnext(new_xml)
    np = Paragraph(new_xml, paragraph._parent)
    if style:
        try:
            np.style = doc.styles[style]
        except KeyError:
            pass
    return np


def add_mixed_runs(paragraph, segments):
    """segments = list of (text, italic_bool)."""
    for text, italic in segments:
        r = paragraph.add_run(text)
        r.italic = italic


# ---------------------------------------------------------------------------
# (c) Parrafo dataset al final de §3.A.1
# ---------------------------------------------------------------------------
idx, p_anchor = find_paragraph_by_substring("Arabidopsis thaliana (Col-0) seeds were sown")
print(f"(c) anchor 3.A.1: parrafo {idx}")

new_para = insert_paragraph_after(p_anchor)
segments_c = [
    (
        "Las 34 muestras del lote 2025017 cubren los cuatro estadios YNG "
        "(juvenil, 21-25 DAS), BOT (",
        False,
    ),
    ("bolting", True),
    (
        ", 28-32 DAS), FLOR (floración, 35-42 DAS) y SEN (senescencia foliar, "
        "45-55 DAS) sobre ",
        False,
    ),
    ("Arabidopsis thaliana", True),
    (
        " Col-0, con n biológicas por estadio YNG = 5, BOT = 11, FLOR = 9 y "
        "SEN = 9. La identificación cronológica de estadios sigue las escalas "
        "de Boyes et al. (2001) y Klepikova et al. (2016). Una muestra "
        "(2025017_15_AT_FLOR) se excluye por sparsity técnica del 80,7 % "
        "(D-002 cap. 4), dejando n = 33 muestras (5+11+8+9) para análisis.",
        False,
    ),
]
add_mixed_runs(new_para, segments_c)
print("(c) parrafo dataset anhadido tras §3.A.1")


# ---------------------------------------------------------------------------
# (d) Tabla nueva con catalogo por familia historica.
# Anclaje: la Tabla 5 (headers "Histone | Region (aa) | ...").
# ---------------------------------------------------------------------------
target_table_xml = None
target_table_idx = -1
for ti, t in enumerate(doc.tables):
    head_text = " ".join(c.text for c in t.rows[0].cells)
    if "Histone" in head_text and "Region" in head_text:
        target_table_xml = t._element
        target_table_idx = ti
        break

print(f"(d) tabla-anchor catalogo: indice tabla {target_table_idx}")

# Parrafo intro tras la tabla 5
intro_p_xml = OxmlElement("w:p")
target_table_xml.addnext(intro_p_xml)
intro_p = Paragraph(intro_p_xml, doc.paragraphs[0]._parent)
intro_p.add_run(
    "El catálogo cuantificado por EpiProfile-PLANTS sobre el lote 2025017 "
    "se distribuye como sigue por familia histónica:"
)

# Tabla nueva: 8 filas (1 header + 7 datos) x 5 cols
new_tbl = doc.add_table(rows=8, cols=5)
new_tbl.style = doc.tables[0].style

headers = ["Familia", "PF post-QC", "Variantes", "Péptidos clave", "Notas"]
filas = [
    ("H1", "6", "H1.1", "H1_104_112", "2 péptidos"),
    ("H2A", "22", "Canónica, W6, W12, X, Xa, Z", "5 péptidos", "variante-específica"),
    ("H2B", "13", "shared, disc-parálogos", "7 péptidos", "K4, K11"),
    ("H3.1", "27", "canónica", "KSAPATGGVKKPHR (H3_27_40)", "A31"),
    ("H3.3", "10", "replication-independent", "KSAPTTGGVKKPHR (H33_27_40)", "T31"),
    ("H4", "9", "canónica",
        "GKGGKGLGKGGAKR (H4_4_17), KVLR (H4_20_23)",
        "K20, acetilaciones"),
    ("TOTAL", "87", "—", "—", "—"),
]

for c_idx, h in enumerate(headers):
    cell = new_tbl.rows[0].cells[c_idx]
    cell.text = ""
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
for r_idx, row in enumerate(filas, start=1):
    is_total = row[0] == "TOTAL"
    for c_idx, val in enumerate(row):
        cell = new_tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        run.bold = is_total

# Mover la tabla nueva (que el add_table coloca al final) para que quede
# justo despues del parrafo introductorio.
intro_p_xml.addnext(new_tbl._element)
print(f"(d) tabla nueva insertada: {len(new_tbl.rows)} filas x {len(new_tbl.columns)} cols")


# ---------------------------------------------------------------------------
# (e) Seccion §3.X "Cuantificación de H3K79"
# ---------------------------------------------------------------------------
heading_p_xml = OxmlElement("w:p")
new_tbl._element.addnext(heading_p_xml)
heading_p = Paragraph(heading_p_xml, doc.paragraphs[0]._parent)

# Intentar estilo Heading 3 / 2 segun disponibilidad
heading_styled = False
for style_name in ("Heading 3", "Heading 2", "Heading 1"):
    try:
        heading_p.style = doc.styles[style_name]
        heading_styled = True
        break
    except KeyError:
        continue
hr = heading_p.add_run("3.X. Cuantificación de H3K79")
hr.bold = True
print(f"(e) heading creado (estilo aplicado: {heading_styled})")

body_p_xml = OxmlElement("w:p")
heading_p_xml.addnext(body_p_xml)
body_p = Paragraph(body_p_xml, doc.paragraphs[0]._parent)
segments_e = [
    (
        "El péptido tríptico EIAQDFKTDLR (H3 residuos 73-83, layout "
        "H3_07_73_83 en EpiProfile-PLANTS) cubre el sitio K79 con cinco "
        "peptidoformas estables en el bundle: K79me1, K79me2, K79me3, "
        "K79ac y la forma sin modificar. K79 es marca ",
        False,
    ),
    ("replication-independent", True),
    (
        " depositada por DOT1L y se asocia a elongación transcripcional en "
        "mamíferos (Steger et al. 2008; Wood, Tellier y Murphy 2018). Su "
        "caracterización en ",
        False,
    ),
    ("Arabidopsis thaliana", True),
    (
        " es escasa: Liu et al. (2010) describe función pero la "
        "cuantificación bottom-up por LC-MS/MS sobre las cuatro formas de "
        "K79 en cuatro estadios ontogénicos es una aportación metodológica "
        "de este bundle.",
        False,
    ),
]
add_mixed_runs(body_p, segments_e)
print("(e) seccion §3.X anhadida tras la tabla nueva")


# ---------------------------------------------------------------------------
# Guardar
# ---------------------------------------------------------------------------
OUT = r"D:/AT_virgen/CHAPTER_METHODS_rev_1604_LV.docx"
doc.save(OUT)
print(f"\nDOCX guardado en: {OUT}")


# ---------------------------------------------------------------------------
# Verificacion
# ---------------------------------------------------------------------------
verify = Document(OUT)
print(f"\nVerificacion: {len(verify.paragraphs)} parrafos, {len(verify.tables)} tablas")

n_old = sum("EpiProfile_PLANTS" in p.text for p in verify.paragraphs)
n_new = sum("EpiProfile-PLANTS" in p.text for p in verify.paragraphs)
print(f"  EpiProfile_PLANTS restantes: {n_old}  (debe ser 0)")
print(f"  EpiProfile-PLANTS presentes:  {n_new}")

patterns_check = [
    "LO PODEMOS REDUCIR",
    "Esta table está bien",
    "Esta figura también está bien",
    "Notas: estas las desarrollaria",
    r"E:\EpiProfile_20_AT",
]
for pat in patterns_check:
    hit = any(pat in p.text for p in verify.paragraphs)
    print(f"  Nota '{pat[:35]}...' presente: {hit} (debe ser False)")

hit_lote = any("lote 2025017" in p.text for p in verify.paragraphs)
hit_k79_h = any("Cuantificación de H3K79" in p.text for p in verify.paragraphs)
hit_k79_body = any("EIAQDFKTDLR" in p.text for p in verify.paragraphs)
hit_catalogo = any(
    "catálogo cuantificado por EpiProfile-PLANTS" in p.text for p in verify.paragraphs
)
print(f"  Parrafo lote 2025017 presente:    {hit_lote}")
print(f"  Heading H3K79 presente:           {hit_k79_h}")
print(f"  Body H3K79 (EIAQDFKTDLR) present: {hit_k79_body}")
print(f"  Intro catalogo presente:          {hit_catalogo}")
